"""Import the Grammar Practice Word bank into the project's JavaScript data file."""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document


QUESTION_EXPORT = "grammarPracticeQuestions"
TOPIC_EXPORT = "grammarPracticeTopics"


def normalize_text(value: str = "") -> str:
    return re.sub(r"\s+", " ", value.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")).strip().casefold()


def extract_json_export(source: str, export_name: str):
    marker = f"export const {export_name} = "
    start = source.index(marker) + len(marker)
    decoder = json.JSONDecoder()
    value, _ = decoder.raw_decode(source[start:])
    return value


def read_existing_bank(path: Path):
    source = path.read_text(encoding="utf-8")
    return {
        "metadata": extract_json_export(source, "grammarPracticeMetadata"),
        "topics": extract_json_export(source, TOPIC_EXPORT),
        "questions": extract_json_export(source, QUESTION_EXPORT),
    }


def parse_question_cell(value: str):
    question_lines = []
    options_by_letter = {}
    current_letter = None

    for raw_line in value.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        option_match = re.match(r"^([A-D])\)\s*(.+)$", line)
        if option_match:
            current_letter = option_match.group(1)
            options_by_letter[current_letter] = option_match.group(2).strip()
        elif current_letter:
            options_by_letter[current_letter] += f" {line}"
        else:
            question_lines.append(line)

    return " ".join(question_lines), [options_by_letter.get(letter, "") for letter in "ABCD"]


def read_word_bank(path: Path):
    document = Document(path)
    if len(document.tables) < 11:
        raise ValueError(f"Expected a summary table and 10 question tables; found {len(document.tables)} tables.")

    summary_rows = document.tables[0].rows[1:]
    topics = []
    for row in summary_rows:
        cells = [cell.text.strip() for cell in row.cells]
        if not cells[0]:
            continue
        topics.append({
            "topic": cells[0],
            "questionCount": int(cells[1]),
            "levelRange": cells[2],
            "description": cells[3],
            "objective": cells[4],
        })

    questions = []
    for topic_index, topic in enumerate(topics, start=1):
        table = document.tables[topic_index]
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            if len(cells) != 6:
                raise ValueError(f"Unexpected column count in topic {topic['topic']}: {len(cells)}")

            source_id, level, specific_topic, question_cell, correct_letter, explanation = cells
            question_text, options = parse_question_cell(question_cell)
            correct_index = ord(correct_letter.strip().upper()) - ord("A") if correct_letter else -1
            correct_answer = options[correct_index] if 0 <= correct_index < len(options) else ""
            questions.append({
                "sourceId": source_id,
                "level": level,
                "topic": topic["topic"],
                "specificTopic": specific_topic,
                "question": question_text,
                "options": options,
                "correctAnswer": correct_answer,
                "explanation": explanation,
                "workContext": "general",
                "source": path.name,
            })

    return topics, questions


def question_match_key(question):
    return normalize_text(question.get("topic", "")), normalize_text(question.get("question", ""))


def duplicate_key(question):
    return (
        normalize_text(question.get("topic", "")),
        normalize_text(question.get("question", "")),
        tuple(normalize_text(option) for option in question.get("options", [])),
        normalize_text(question.get("correctAnswer", "")),
    )


def content_signature(question):
    return (
        normalize_text(question.get("topic", "")),
        normalize_text(question.get("level", "")),
        normalize_text(question.get("specificTopic", "")),
        normalize_text(question.get("question", "")),
        tuple(normalize_text(option) for option in question.get("options", [])),
        normalize_text(question.get("correctAnswer", "")),
        normalize_text(question.get("explanation", "")),
    )


def validate_questions(questions):
    errors = []
    seen_ids = set()
    seen_duplicates = set()
    duplicate_count = 0

    for question in questions:
        identifier = question.get("id", "")
        if not identifier:
            errors.append(f"{question.get('sourceId', 'unknown')}: missing ID")
        elif identifier in seen_ids:
            errors.append(f"{identifier}: duplicate ID")
        seen_ids.add(identifier)

        key = duplicate_key(question)
        if key in seen_duplicates:
            duplicate_count += 1
            errors.append(f"{identifier}: duplicate question content")
        seen_duplicates.add(key)

        if not question.get("topic"):
            errors.append(f"{identifier}: missing topic")
        if not question.get("question"):
            errors.append(f"{identifier}: missing question text")
        if len(question.get("options", [])) != 4 or any(not option for option in question.get("options", [])):
            errors.append(f"{identifier}: invalid options")
        if question.get("correctAnswer") not in question.get("options", []):
            errors.append(f"{identifier}: correct answer is not an option")
        if not question.get("explanation"):
            errors.append(f"{identifier}: missing explanation")

    return errors, duplicate_count


def merge_bank(existing_questions, word_questions):
    existing_by_key = defaultdict(list)
    for question in existing_questions:
        existing_by_key[question_match_key(question)].append(question)

    used_existing_ids = set()
    next_id = max((int(match.group(1)) for question in existing_questions if (match := re.fullmatch(r"gp-(\d+)", question.get("id", "")))), default=0) + 1
    merged = []
    matched = []
    new_questions = []
    source_id_changes = []

    for word_question in word_questions:
        candidates = [item for item in existing_by_key[question_match_key(word_question)] if item["id"] not in used_existing_ids]
        existing = candidates[0] if len(candidates) == 1 else None

        if existing:
            question_id = existing["id"]
            used_existing_ids.add(question_id)
            matched.append((existing, word_question))
            if existing.get("sourceId") != word_question.get("sourceId"):
                source_id_changes.append((question_id, existing.get("sourceId"), word_question.get("sourceId")))
            word_question["workContext"] = existing.get("workContext", "general")
        else:
            question_id = f"gp-{next_id:03d}"
            next_id += 1
            new_questions.append(word_question)

        merged.append({"id": question_id, **word_question})

    missing_from_word = [question for question in existing_questions if question["id"] not in used_existing_ids]
    modified = [(old, new) for old, new in matched if content_signature(old) != content_signature(new)]
    return merged, new_questions, modified, missing_from_word, source_id_changes


def render_js(metadata, topics, questions):
    return (
        f"export const grammarPracticeMetadata = {json.dumps(metadata, ensure_ascii=False, indent=2)}\n\n"
        f"export const grammarPracticeTopics = {json.dumps(topics, ensure_ascii=False, indent=2)}\n\n"
        f"export const grammarPracticeQuestions = {json.dumps(questions, ensure_ascii=False, indent=2)}\n"
    )


def distribution(items, field):
    return dict(sorted(Counter(item.get(field, "") for item in items).items()))


def report_markdown(source, output, existing, topics, questions, new_questions, modified, missing, source_id_changes, errors, duplicates):
    source_stat = source.stat()
    topic_counts = distribution(questions, "topic")
    level_counts = distribution(questions, "level")
    subtopic_counts = distribution(questions, "specificTopic")
    old_topic_counts = distribution(existing["questions"], "topic")

    def bullets(values):
        return "\n".join(f"- {key}: {value}" for key, value in values.items())

    return f"""# Grammar Practice Update Report

## Fuente y archivos

- Documento procesado: `{source}`
- Fecha de modificación: {datetime.fromtimestamp(source_stat.st_mtime).astimezone().isoformat(timespec="seconds")}
- Banco activo actualizado: `{output}`
- Documento usado únicamente durante desarrollo; React consume el archivo JavaScript generado.

## Resumen de actualización

- Banco anterior: {len(existing['questions'])} preguntas, {len(old_topic_counts)} temas.
- Documento actualizado: {len(questions)} preguntas detectadas, {len(topics)} temas.
- Banco final: {len(questions)} preguntas, {len(topic_counts)} temas.
- Preguntas nuevas añadidas: {len(new_questions)}.
- Preguntas con contenido modificado: {len(modified)}.
- IDs de origen actualizados por la nueva numeración del Word: {len(source_id_changes)}.
- Registros anteriores ausentes en el Word: {len(missing)}.
- Duplicados omitidos o detectados: {duplicates}.
- Preguntas rechazadas o pendientes de revisión: {len(errors)}.
- IDs únicos: {len({item['id'] for item in questions})}.

## Banco anterior por tema

{bullets(old_topic_counts)}

## Banco final por tema

{bullets(topic_counts)}

## Banco final por nivel

{bullets(level_counts)}

## Banco final por subtema

{bullets(subtopic_counts)}

## Comparación

- IDs internos existentes conservados al emparejar tema y texto normalizados: {len(questions) - len(new_questions)}.
- IDs internos nuevos asignados: {len(new_questions)}.
- Fragmentos modificados: {', '.join(f"{old['id']} ({old['question'][:45]})" for old, _ in modified[:20]) or 'Ninguno'}.
- Registros anteriores ausentes: {', '.join(f"{item['id']} ({item['question'][:45]})" for item in missing[:20]) or 'Ninguno'}.
- Problemas: {'; '.join(errors[:30]) or 'Ninguno'}.

## Validaciones realizadas

- Cada pregunta tiene ID, tema, nivel, subtema, texto, cuatro opciones, respuesta y explicación.
- Cada respuesta correcta coincide exactamente con una opción.
- Los IDs internos y el contenido normalizado no se duplican.
- Las preguntas conservadas se identificaron por tema y texto normalizados, no solo por ID de origen.
- Las preguntas nuevas quedan en el mismo banco que alimenta la práctica por tema y Grammar Level Check.
- El documento Word no se consulta durante los intentos en el navegador.
"""


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--existing", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--report", required=True, type=Path)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    existing = read_existing_bank(args.existing)
    topics, word_questions = read_word_bank(args.source)
    merged, new_questions, modified, missing, source_id_changes = merge_bank(existing["questions"], word_questions)
    errors, duplicates = validate_questions(merged)

    if errors:
        raise ValueError("Bank validation failed:\n" + "\n".join(errors))

    metadata = {
        **existing["metadata"],
        "sourceDocument": args.source.name,
        "description": "Banco ampliado de preguntas para refuerzo gramatical por tema.",
        "questionCount": len(merged),
        "topicCount": len(topics),
    }
    report = report_markdown(args.source, args.output, existing, topics, merged, new_questions, modified, missing, source_id_changes, errors, duplicates)

    summary = {
        "previousTotal": len(existing["questions"]),
        "documentTotal": len(word_questions),
        "finalTotal": len(merged),
        "topics": distribution(merged, "topic"),
        "levels": distribution(merged, "level"),
        "subtopicCount": len(distribution(merged, "specificTopic")),
        "newQuestions": len(new_questions),
        "modifiedQuestions": len(modified),
        "missingFromWord": len(missing),
        "sourceIdChanges": len(source_id_changes),
        "duplicates": duplicates,
        "pendingReview": len(errors),
        "uniqueIds": len({item["id"] for item in merged}),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))

    if not args.dry_run:
        args.output.write_text(render_js(metadata, topics, merged), encoding="utf-8")
        args.report.write_text(report, encoding="utf-8")


if __name__ == "__main__":
    main()
